import streamlit as st
import pandas as pd
import os
from io import BytesIO
# import docx
from docx import Document

st.set_page_config(page_title="Data Sweeper", layout='wide')
st.title("Data Sweeper")
st.write("Transform your files between CSV, Excel, and DOCX formats with built-in data cleaning and visualization!")

# Function to Read DOCX and Convert to DataFrame
def read_docx(file):
    doc = Document(file)
    text_data = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_data.append([para.text.strip()])
    
    if not text_data:
        return None
    return pd.DataFrame(text_data, columns=["Content"])

# Function to Convert DataFrame to DOCX
def convert_to_docx(df, file_name):
    doc = Document()
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    
    # Add Headers
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = col_name
    
    # Add Data
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer, f"{file_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Function to Convert DOCX to Excel
def convert_docx_to_excel(file):
    df = read_docx(file)
    if df is None:
        return None, None, None

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Sheet1")
    buffer.seek(0)

    return buffer, "converted.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

# Function to Convert DOCX to CSV
def convert_docx_to_csv(file):
    df = read_docx(file)
    if df is None:
        return None, None, None

    buffer = BytesIO()
    df.to_csv(buffer, index=False)
    buffer.seek(0)

    return buffer, "converted.csv", "text/csv"

# Upload Files
uploaded_files = st.file_uploader("Upload your files (CSV, Excel, or DOCX):", type=["csv", "xlsx", "docx"], accept_multiple_files=True)

if uploaded_files:
    for file in uploaded_files:
        file_ext = os.path.splitext(file.name)[-1].lower()

        # Read Files
        if file_ext == ".csv":
            df = pd.read_csv(file)
        elif file_ext == ".xlsx":
            df = pd.read_excel(file, engine="openpyxl")
        elif file_ext == ".docx":
            df = read_docx(file)
            if df is None:
                st.error(f"Unable to process DOCX file: {file.name}")
                continue
        else:
            st.error(f"Unsupported file type: {file_ext}")
            continue

        st.write(f'**File Name:** {file.name}')
        st.write(f"**File Size:** {file.size / 1024:.2f} KB")

        # Preview Data
        st.subheader("Preview of the Data")
        st.dataframe(df.head())

        # Data Cleaning Options
        st.subheader("Data Cleaning Options")
        if st.checkbox(f"Clean Data for {file.name}"):
            col1, col2 = st.columns(2)

            with col1:
                if st.button(f"Remove Duplicates from {file.name}"):
                    df.drop_duplicates(inplace=True)
                    st.write("Duplicates Removed!")

            with col2:
                if st.button(f"Fill Missing Values for {file.name}"):
                    numeric_cols = df.select_dtypes(include=["number"]).columns
                    df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].mean())
                    st.write("Missing values have been filled!")

        # Column Selection
        st.subheader("Select Columns to Keep")
        columns = st.multiselect(f"Choose Columns for {file.name}", df.columns, default=df.columns)
        df = df[columns]

        # Data Visualization
        st.subheader("Data Visualization")
        if st.checkbox(f"Show Visualization for {file.name}"):
            st.bar_chart(df.select_dtypes(include='number').iloc[:, :2])

        # Conversion Options
        st.subheader("Conversion Options")
        conversion_type = st.radio(f"Convert {file.name} to:", ["CSV", "Excel", "DOCX"], key=file.name)

        if st.button(f"Convert {file.name}"):
            buffer = BytesIO()
            new_file_name = file.name.rsplit(".", 1)[0]

            if conversion_type == "CSV":
                df.to_csv(buffer, index=False)
                buffer.seek(0)
                file_name = f"{new_file_name}.csv"
                mime_type = "text/csv"

            elif conversion_type == "Excel":
                with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
                    df.to_excel(writer, index=False, sheet_name="Sheet1")
                buffer.seek(0)
                file_name = f"{new_file_name}.xlsx"
                mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

            elif conversion_type == "DOCX":
                buffer, file_name, mime_type = convert_to_docx(df, new_file_name)

            # Download Button
            st.download_button(
                label=f"Download {file_name}",
                data=buffer,
                file_name=file_name,
                mime=mime_type
            )

    st.success("All files processed successfully! 🎉")
